# -*- coding: utf-8 -*-
"""Convert the translated Markdown book to a formatted Word document."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

MD_PATH = Path(r"C:\Users\bisu5\Desktop\南亚研究skills\《命运的四颗星》纳拉瓦内回忆录_中文译本.md")
DOCX_PATH = Path(r"C:\Users\bisu5\Desktop\南亚研究skills\《命运的四颗星》纳拉瓦内回忆录_中文译本.docx")

def set_cell_font(run, name_cn='宋体', name_en='Times New Roman', size=12, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.name = name_en
    run.font.bold = bold
    # Set East Asian font
    from docx.oxml.ns import qn
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name_cn)
    if color:
        run.font.color.rgb = RGBColor(*color)

def create_word_doc(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Page setup - A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # Default paragraph style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    from docx.oxml.ns import qn
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Skip horizontal rules
        if stripped == '---':
            i += 1
            continue

        # Blockquote (translator's note)
        if stripped.startswith('> '):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(text)
            set_cell_font(run, '楷体', 'Times New Roman', 10.5, color=(100, 100, 100))
            i += 1
            continue

        # H1 - Book title
        if stripped.startswith('# ') and not stripped.startswith('## '):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(36)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(text)
            set_cell_font(run, '黑体', 'Times New Roman', 22, bold=True)
            i += 1
            continue

        # H2 - Parts and major sections
        if stripped.startswith('## ') and not stripped.startswith('### '):
            text = stripped[3:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            set_cell_font(run, '黑体', 'Times New Roman', 16, bold=True)
            # Add a thin line below part headings
            i += 1
            continue

        # H3 - Chapters
        if stripped.startswith('### '):
            text = stripped[4:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            set_cell_font(run, '黑体', 'Times New Roman', 14, bold=True)
            i += 1
            continue

        # Bold text line (like subtitle)
        if stripped.startswith('**') and stripped.endswith('**'):
            text = stripped[2:-2]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            set_cell_font(run, '宋体', 'Times New Roman', 12, bold=True)
            i += 1
            continue

        # Italic text line (like author info)
        if stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**'):
            text = stripped[1:-1]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            set_cell_font(run, '楷体', 'Times New Roman', 12)
            run.font.italic = True
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)  # 2-char indent
        # Handle inline formatting
        text = stripped
        # Simple: just add as plain text (inline bold/italic rare in body)
        run = p.add_run(text)
        set_cell_font(run, '宋体', 'Times New Roman', 12)

        i += 1

    doc.save(str(docx_path))
    print(f"Word document saved: {docx_path}")
    print(f"File size: {docx_path.stat().st_size} bytes")

if __name__ == '__main__':
    create_word_doc(MD_PATH, DOCX_PATH)
